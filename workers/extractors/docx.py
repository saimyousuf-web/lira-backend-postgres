import uuid
import json
import zipfile
from datetime import datetime, timezone
from io import BytesIO

import pytesseract
from PIL import Image
from docx import Document as DocxDocument
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter

from fastapi import APIRouter, HTTPException
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.exc import IntegrityError, SQLAlchemyError
from sqlalchemy import select

from models.chunks import Chunk
from models.module import Module
from models.course_module import CourseModule
from workers.vector import bedrock_client, normalize_vector
from qdrant_client.models import PointStruct
from workers.qdrnt_vector import get_qdrant_client, QDRANT_COLLECTION

router = APIRouter()


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def extract_docx_pages(doc_bytes: bytes) -> dict:
    """
    Extract full text from a DOCX file (digital text + OCR on embedded images).
    Returns: {"pages": [{"text": "<combined text>"}]}
    """
    # --- Digital text ---
    digital_parts = []
    try:
        doc = DocxDocument(BytesIO(doc_bytes))
        for p in doc.paragraphs:
            t = p.text.strip()
            if t:
                digital_parts.append(t)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    t = cell.text.strip()
                    if t:
                        digital_parts.append(t)
    except Exception as e:
        print("Error extracting DOCX digital text:", e)

    digital_text = "\n".join(digital_parts)

    # --- OCR on embedded images ---
    ocr_text = ""
    try:
        with zipfile.ZipFile(BytesIO(doc_bytes)) as z:
            for filename in z.namelist():
                if filename.startswith("word/media/"):
                    try:
                        img = Image.open(BytesIO(z.read(filename)))
                        text = pytesseract.image_to_string(img).strip()
                        if text:
                            ocr_text += "\n" + text
                    except Exception as e:
                        print(f"OCR failed for {filename}: {e}")
    except Exception as e:
        print("Error reading DOCX zip for images:", e)

    return {"pages": [{"text": (digital_text + "\n" + ocr_text).strip()}]}


def _build_chunks(docx_json: dict) -> list[dict]:
    """Split extracted text into overlapping chunks."""
    pages = docx_json.get("pages", [])
    if not pages:
        return []

    page_text = pages[0].get("text", "").strip()
    if not page_text:
        return []

    splitter = RecursiveCharacterTextSplitter(chunk_size=1200, chunk_overlap=150)
    raw_chunks = splitter.split_documents([Document(page_content=page_text)])

    return [
        {
            "chunk_id": uuid.uuid4(),
            "text": chunk.page_content,
            "chunk_index": idx,
        }
        for idx, chunk in enumerate(raw_chunks, start=1)
    ]


def _embed_and_upsert_qdrant(
    chunk_id: uuid.UUID,
    text: str,
    organization_id: str,
    course_id: uuid.UUID,
    course_name: str,
    module_id: uuid.UUID,
) -> tuple[str, int]:
    """
    Embed a single chunk with Amazon Titan and upsert into Qdrant.
    Returns (embedding_status, embedding_dim).
    """
    try:
        response = bedrock_client.invoke_model(
            modelId="amazon.titan-embed-text-v2:0",
            contentType="application/json",
            body=json.dumps({"inputText": text}),
        )
        body = json.loads(response["body"].read())
        embedding = body.get("embedding")

        if not (isinstance(embedding, list) and len(embedding) == 1024):
            return "failed", 0
        print(get_qdrant_client().collection_exists(QDRANT_COLLECTION))
        normalized = normalize_vector(embedding)
        get_qdrant_client().upsert(
        collection_name=QDRANT_COLLECTION,
        points=[
        PointStruct(
            id=str(chunk_id),     
            vector=normalized,
            payload={
                "chunk_id":       str(chunk_id),
                "organization_id": organization_id,
                "course_id":      str(course_id),
                "course_name":    course_name,
                "module_id":      str(module_id),
                "slide_index": 1,  
             },
            )
         ],
        )
        return "completed", len(embedding)

    except Exception as e:
        print(f"Embedding failed for chunk {chunk_id}: {e}")
        return "failed", 0



async def chunk_and_embed_docx(
    docx_bytes: bytes,
    ctx_orgid: str,
    ctx_cid: str,
    course_name: str,
    ctx_moid: str,
    db : AsyncSession,
):
    """
    Chunk and embed a DOCX module document.
    """
    # Hardcoded acting user — replace with real auth dependency
    user_id = uuid.UUID("6418e458-50a1-70fe-9d3e-b52f5d2df57c")

    try:
        course_id = ctx_cid
        module_id = ctx_moid
    except ValueError:
        raise HTTPException(status_code=400, detail="Invalid UUID format in path")

    # ------------------------------------------------------------------
    # 1. Validate Module + CourseModule exist
    # ------------------------------------------------------------------
    try:
        module_result = await db.execute(select(Module).where(Module.id == module_id))
        module: Module | None = module_result.scalar_one_or_none()
        if not module:
            raise HTTPException(status_code=404, detail="Module not found")

        cm_result = await db.execute(
            select(CourseModule).where(
                CourseModule.cid == course_id,
                CourseModule.moid == module_id,
            )
        )
        course_module: CourseModule | None = cm_result.scalar_one_or_none()
        if not course_module:
            raise HTTPException(status_code=404, detail="CourseModule association not found")

    except HTTPException:
        raise
    except SQLAlchemyError as e:
        raise HTTPException(status_code=500, detail=f"Database error: {e}")


    # 3. Extract text
    docx_json = extract_docx_pages(docx_bytes)
    raw_chunks = _build_chunks(docx_json)

    if not raw_chunks:
        raise HTTPException(status_code=422, detail="No text could be extracted from document")

    # ------------------------------------------------------------------
    # 4. Embed + persist each chunk
    # ------------------------------------------------------------------
    course_name = module.nm   # fallback; pass real course name if available

    for chunk_data in raw_chunks:
        chunk_id: uuid.UUID = chunk_data["chunk_id"]
        text: str = chunk_data["text"]

        embedding_status, _ = _embed_and_upsert_qdrant(
            chunk_id=chunk_id,
            text=text,
            organization_id=ctx_orgid,
            course_id=course_id,
            course_name=course_name,
            module_id=module_id,
        )

        if embedding_status == "failed":
            print(embedding_status, 'embedding_statusembedding_statusembedding_status')
            print(f"Skipping DB insert for chunk {chunk_id} due to embedding failure")
            continue

        chunk_row = Chunk(
            id=chunk_id,
            moid=module_id,
            cid=course_id,
            txt=text,
            crtby=user_id,
            updby=user_id,
        )
        db.add(chunk_row)


    # 5. Mark Module + CourseModule as vectorised / processed
    now = datetime.now(timezone.utc)

    module.isvec = True
    module.sts = "APPROVED"
    module.updat = now
    module.updby = user_id

    course_module.updat = now
    course_module.updby = user_id


    # 6. Commit everything in one transaction
    try:
        await db.commit()
    except IntegrityError as e:
        await db.rollback()
        raise HTTPException(status_code=500, detail=f"Integrity error: {e.orig}")
    except SQLAlchemyError as e:
        await db.rollback()
        raise HTTPException(status_code=500, detail=f"Database error: {e}")

    print(f"DOCX chunking + embedding complete. {len(raw_chunks)} chunks saved.")

    return {
        "status": "success",
        "status_code": 200,
        "message": "Document chunked and embedded successfully",
        "chunk_count": len(raw_chunks),
    }